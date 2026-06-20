#!/usr/bin/env python3
"""
report_docx.py — Export Word (.docx) du rapport, en parallèle du PDF.

Produit un .docx éditable (le client peut ajouter son logo, modifier le nom de
société, etc.) avec le même thème PenTest Hub : titres rouges, tableaux de
sévérité colorés, blocs d'évidence en monospace.

Nécessite python-docx (`pip install python-docx`). Si le module est absent,
l'appelant ignore proprement cette sortie.

Le mapping (réinjection des vraies valeurs) est appliqué AVANT par l'appelant :
ce module reçoit déjà le Markdown final.
"""
from __future__ import annotations

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False

RED   = "E63946"
DARK  = "2B2D42"
WHITE = "FFFFFF"
SEV_FILL = {
    "CRITICAL": "E63946", "HIGH": "F4845F", "MEDIUM": "F4D35E",
    "LOW": "06D6A0", "INFO": "118AB2",
}


def _shade(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_text(cell, text, *, bold=False, color=None, size=9, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_md_table(doc, lines):
    rows = []
    for ln in lines:
        if re.match(r"^\s*\|[-| :]+\|\s*$", ln):
            continue
        rows.append([c.strip() for c in ln.strip().strip("|").split("|")])
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            key = val.strip().upper()
            if ri == 0:
                _set_cell_text(cell, val, bold=True, color=WHITE, size=9)
                _shade(cell, DARK)
            elif key in SEV_FILL and ci == ncol - 1:
                txtcol = WHITE if key in ("CRITICAL", "HIGH", "INFO") else DARK
                _set_cell_text(cell, key, bold=True, color=txtcol, size=9, align="center")
                _shade(cell, SEV_FILL[key])
            else:
                _set_cell_text(cell, val, size=9)
    doc.add_paragraph()


def _heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18); run.font.color.rgb = RGBColor.from_string(DARK)
    elif level == 2:
        run.font.size = Pt(14); run.font.color.rgb = RGBColor.from_string(RED)
    elif level == 3:
        run.font.size = Pt(11.5); run.font.color.rgb = RGBColor.from_string(DARK)
    else:
        run.font.size = Pt(10.5); run.font.color.rgb = RGBColor.from_string(DARK)
    return p


def convert_md_to_docx(md_text: str, docx_path: str,
                       classification: str = "CONFIDENTIAL") -> bool:
    """Génère un .docx depuis le Markdown final. Retourne True si OK."""
    if not _DOCX_OK:
        return False

    doc = Document()
    # Marges
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)
        # Pied de page avec classification
        footer = s.footer.paragraphs[0]
        footer.text = f"{classification}  ·  RMW-PenStrike  ·  www.rmw-penstrike.com"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in footer.runs:
            r.font.size = Pt(7.5); r.font.color.rgb = RGBColor.from_string("6C757D")

    # Bandeau de titre
    band = doc.add_paragraph()
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = band.add_run(f"PENETRATION TESTING REPORT  —  {classification}")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(RED)
    doc.add_paragraph()

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            # léger fond gris via shading du paragraphe
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            code_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                flush_code(); in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        m = re.match(r"^(#{1,4})\s+(.+)", line)
        if m:
            _heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1; continue

        if line.strip().startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i]); i += 1
            _add_md_table(doc, tbl); continue

        if re.match(r"^---+\s*$", line.strip()):
            i += 1; continue

        m = re.match(r"^(\s*)[-*+]\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, m.group(2)); i += 1; continue

        if not line.strip():
            i += 1; continue

        p = doc.add_paragraph()
        _add_inline(p, line)
        i += 1

    flush_code()
    try:
        doc.save(docx_path)
        return True
    except Exception:
        return False


def _add_inline(paragraph, text):
    """Gère **bold** et `code` inline."""
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); run.font.name = "Courier New"
        else:
            paragraph.add_run(part)
        for r in paragraph.runs:
            if r.font.size is None:
                r.font.size = Pt(9.5)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python3 report_docx.py RAPPORT.md [SORTIE.docx]")
        sys.exit(1)
    md = Path(sys.argv[1]).read_text(encoding="utf-8", errors="replace")
    out = sys.argv[2] if len(sys.argv) > 2 else "rapport.docx"
    ok = convert_md_to_docx(md, out)
    print("OK ->", out if ok else "ÉCHEC (python-docx manquant ?)")
